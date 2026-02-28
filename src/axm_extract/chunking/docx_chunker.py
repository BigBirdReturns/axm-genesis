from __future__ import annotations

from dataclasses import asdict
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document  # python-docx

from axm_extract.models.types import Chunk, Locator, TableChunk, TextSpan
from axm_extract.chunking.util import make_content_anchored_chunk_id, norm_ws, sha256_hex

def _doc_id_from_path(path: Path) -> str:
    # Deterministic ID for v1: content hash of raw bytes
    data = path.read_bytes()
    import hashlib
    return "sha256:" + hashlib.sha256(data).hexdigest()

def chunk_docx(path: Path) -> Tuple[str, str, List[object]]:
    doc = Document(str(path))
    doc_id = _doc_id_from_path(path)

    extracted_parts: List[str] = []
    chunks: List[object] = []

    cursor = 0

    # Paragraphs as prose chunks
    para_index = 0
    for p in doc.paragraphs:
        text = norm_ws(p.text or "")
        if not text:
            para_index += 1
            continue

        # Append with newline separator
        start = cursor
        extracted_parts.append(text)
        cursor += len(text)
        end = cursor
        extracted_parts.append("\n\n")
        cursor += 2

        locator = Locator(kind="docx", paragraph_index=para_index, run_index=None)
        locator_key = f"p:{para_index}"
        chunk_id = make_content_anchored_chunk_id(doc_id, "prose", locator_key, text)

        chunks.append(
            Chunk(
                chunk_id=chunk_id,
                chunk_type="prose",
                locator=locator,
                text_span=TextSpan(artifact="extracted_text", start=start, end=end),
                text=text,
            )
        )

        para_index += 1

    # Tables as table chunks
    # python-docx tables are not interleaved with paragraphs in order here.
    # v1: emit table chunks after prose chunks. Locator uses block_id = "table:{i}".
    table_index = 0
    for t in doc.tables:
        headers: List[str] = []
        rows: List[List[str]] = []
        cell_spans: Dict[str, TextSpan] = {}

        # Simple heuristic: first row is headers if it looks header-like
        raw = [[norm_ws(c.text or "") for c in row.cells] for row in t.rows]
        raw = [r for r in raw if any(x for x in r)]
        if not raw:
            table_index += 1
            continue

        if len(raw) >= 2:
            headers = raw[0]
            rows = raw[1:]
        else:
            headers = raw[0]
            rows = []

        # Build flattened table text and compute spans per cell.
        table_lines: List[str] = []
        # Header line
        if headers:
            table_lines.append("\t".join(headers))

        # Data lines
        for r in rows:
            table_lines.append("\t".join(r))

        table_text = "\n".join(table_lines)
        if not table_text:
            table_index += 1
            continue

        # Append table text to extracted_text
        start_table = cursor
        extracted_parts.append(table_text)
        cursor += len(table_text)
        end_table = cursor
        extracted_parts.append("\n\n")
        cursor += 2

        # Compute per-cell spans by reconstructing offsets in the flattened table_text
        # This is deterministic because we control serialization (tabs, newlines).
        # Header row spans
        offset = 0
        def _row_spans(row_vals: List[str], row_idx: int):
            nonlocal offset
            for col_idx, val in enumerate(row_vals):
                v = val or ""
                cell_start = offset
                cell_end = offset + len(v)
                cell_spans[f"{row_idx}:{col_idx}"] = TextSpan(
                    artifact="extracted_text",
                    start=start_table + cell_start,
                    end=start_table + cell_end,
                )
                offset = cell_end
                if col_idx < len(row_vals) - 1:
                    offset += 1  # tab
            offset += 1  # newline

        if headers:
            _row_spans(headers, 0)

        for i, r in enumerate(rows, start=1):
            _row_spans(r, i)

        locator = Locator(kind="docx", paragraph_index=None, run_index=None, table_id=f"table:{table_index}")
        locator_key = f"table:{table_index}"
        chunk_id = make_content_anchored_chunk_id(doc_id, "table", locator_key, table_text)

        chunks.append(
            TableChunk(
                chunk_id=chunk_id,
                chunk_type="table",
                locator=locator,
                text=table_text,
                headers=headers,
                rows=rows,
                cell_spans=cell_spans,
                text_span=TextSpan(artifact="extracted_text", start=start_table, end=end_table),
            )
        )

        table_index += 1

    extracted_text = "".join(extracted_parts)
    return doc_id, extracted_text, chunks
